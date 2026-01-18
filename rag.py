import streamlit as st
import os
from pathlib import Path
from dotenv import load_dotenv
from langchain_groq import ChatGroq
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_community.vectorstores import FAISS
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.messages import HumanMessage, AIMessage
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser
import tempfile
from fpdf import FPDF
from datetime import datetime
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Load environment variables
load_dotenv()

# Page configuration
st.set_page_config(
    page_title="Study-Buddy",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Apple-inspired CSS
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
    
    * {
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
    }
    
    .main {
        background: linear-gradient(135deg, #0a0a0a 0%, #1a1a1a 100%);
        padding: 0;
    }
    
    .stApp {
        background: transparent;
    }
    
    /* Hero Section */
    .hero-section {
        text-align: center;
        padding: 8rem 2rem 4rem 2rem;
        background: linear-gradient(180deg, rgba(0,0,0,0) 0%, rgba(0,122,255,0.1) 100%);
        border-radius: 0 0 50px 50px;
        margin-bottom: 3rem;
    }
    
    .hero-title {
        font-size: 4.5rem;
        font-weight: 700;
        background: linear-gradient(135deg, #ffffff 0%, #007AFF 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        margin-bottom: 1rem;
        letter-spacing: -0.02em;
        line-height: 1.1;
    }
    
    .hero-subtitle {
        font-size: 1.5rem;
        color: #86868b;
        font-weight: 400;
        margin-bottom: 3rem;
        letter-spacing: -0.01em;
    }
    
    /* Feature Cards */
    .feature-card {
        background: rgba(255, 255, 255, 0.05);
        backdrop-filter: blur(20px);
        -webkit-backdrop-filter: blur(20px);
        border: 1px solid rgba(255, 255, 255, 0.1);
        border-radius: 24px;
        padding: 2.5rem;
        margin-bottom: 1.5rem;
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
        box-shadow: 0 4px 24px rgba(0, 0, 0, 0.3);
    }
    
    .feature-card:hover {
        transform: translateY(-8px);
        box-shadow: 0 12px 40px rgba(0, 122, 255, 0.3);
        border-color: rgba(0, 122, 255, 0.3);
    }
    
    .feature-icon {
        font-size: 3rem;
        margin-bottom: 1rem;
        display: block;
    }
    
    .feature-title {
        font-size: 1.5rem;
        font-weight: 600;
        color: #ffffff;
        margin-bottom: 0.75rem;
    }
    
    .feature-desc {
        font-size: 1.05rem;
        color: #86868b;
        line-height: 1.6;
    }
    
    /* Chat Interface */
    .chat-container {
        background: rgba(255, 255, 255, 0.03);
        backdrop-filter: blur(30px);
        border-radius: 30px;
        padding: 2rem;
        margin: 2rem auto;
        max-width: 900px;
        border: 1px solid rgba(255, 255, 255, 0.08);
        box-shadow: 0 8px 32px rgba(0, 0, 0, 0.4);
    }
    
    .stChatMessage {
        background: rgba(255, 255, 255, 0.05) !important;
        border: 1px solid rgba(255, 255, 255, 0.1) !important;
        border-radius: 18px !important;
        padding: 1.25rem !important;
        margin-bottom: 1rem !important;
        backdrop-filter: blur(10px);
    }
    
    .stChatMessage[data-testid="user-message"] {
        background: linear-gradient(135deg, rgba(0, 122, 255, 0.15) 0%, rgba(0, 122, 255, 0.05) 100%) !important;
        border-color: rgba(0, 122, 255, 0.3) !important;
    }
    
    .stChatMessage[data-testid="assistant-message"] {
        background: rgba(255, 255, 255, 0.03) !important;
    }
    
    /* Buttons */
    .stButton > button {
        background: linear-gradient(135deg, #007AFF 0%, #0051D5 100%);
        color: white;
        border: none;
        border-radius: 14px;
        padding: 0.875rem 2rem;
        font-weight: 600;
        font-size: 1.05rem;
        letter-spacing: 0.01em;
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
        box-shadow: 0 4px 16px rgba(0, 122, 255, 0.4);
    }
    
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 8px 24px rgba(0, 122, 255, 0.5);
        background: linear-gradient(135deg, #0051D5 0%, #003DA5 100%);
    }
    
    /* Input Fields */
    .stTextInput > div > div > input,
    .stChatInput > div > div > input {
        background: rgba(255, 255, 255, 0.05) !important;
        border: 1px solid rgba(255, 255, 255, 0.15) !important;
        border-radius: 12px !important;
        color: #ffffff !important;
        padding: 0.875rem 1.25rem !important;
        font-size: 1rem !important;
        backdrop-filter: blur(10px);
    }
    
    .stTextInput > div > div > input:focus,
    .stChatInput > div > div > input:focus {
        border-color: rgba(0, 122, 255, 0.5) !important;
        box-shadow: 0 0 0 3px rgba(0, 122, 255, 0.1) !important;
    }
    
    /* Upload Section */
    .upload-section {
        background: linear-gradient(135deg, rgba(0, 122, 255, 0.1) 0%, rgba(88, 86, 214, 0.1) 100%);
        border: 2px dashed rgba(0, 122, 255, 0.3);
        border-radius: 24px;
        padding: 3rem;
        text-align: center;
        margin: 2rem 0;
        transition: all 0.3s ease;
    }
    
    .upload-section:hover {
        border-color: rgba(0, 122, 255, 0.6);
        background: linear-gradient(135deg, rgba(0, 122, 255, 0.15) 0%, rgba(88, 86, 214, 0.15) 100%);
    }
    
    /* Status Badge */
    .status-badge {
        display: inline-block;
        padding: 0.5rem 1.25rem;
        border-radius: 20px;
        font-size: 0.95rem;
        font-weight: 600;
        letter-spacing: 0.02em;
    }
    
    .status-active {
        background: rgba(52, 199, 89, 0.15);
        color: #34C759;
        border: 1px solid rgba(52, 199, 89, 0.3);
    }
    
    .status-inactive {
        background: rgba(255, 59, 48, 0.15);
        color: #FF3B30;
        border: 1px solid rgba(255, 59, 48, 0.3);
    }
    
    /* Expandable Sections */
    .streamlit-expanderHeader {
        background: rgba(255, 255, 255, 0.05) !important;
        border-radius: 12px !important;
        color: #ffffff !important;
        font-weight: 500 !important;
        border: 1px solid rgba(255, 255, 255, 0.1) !important;
    }
    
    .streamlit-expanderContent {
        background: rgba(255, 255, 255, 0.02) !important;
        border-radius: 0 0 12px 12px !important;
        border: 1px solid rgba(255, 255, 255, 0.1) !important;
        border-top: none !important;
    }
    
    /* File Uploader */
    .stFileUploader {
        background: transparent !important;
    }
    
    .stFileUploader > div {
        background: rgba(255, 255, 255, 0.05) !important;
        border: 2px dashed rgba(255, 255, 255, 0.2) !important;
        border-radius: 16px !important;
        padding: 2rem !important;
    }
    
    /* Download Button */
    .stDownloadButton > button {
        background: linear-gradient(135deg, #34C759 0%, #30B350 100%) !important;
        box-shadow: 0 4px 16px rgba(52, 199, 89, 0.4) !important;
    }
    
    .stDownloadButton > button:hover {
        background: linear-gradient(135deg, #30B350 0%, #28A745 100%) !important;
        box-shadow: 0 8px 24px rgba(52, 199, 89, 0.5) !important;
    }
    
    /* Metrics */
    .stMetric {
        background: rgba(255, 255, 255, 0.05);
        padding: 1.5rem;
        border-radius: 16px;
        border: 1px solid rgba(255, 255, 255, 0.1);
    }
    
    .stMetric label {
        color: #86868b !important;
        font-size: 0.95rem !important;
    }
    
    .stMetric [data-testid="stMetricValue"] {
        color: #ffffff !important;
        font-size: 1.75rem !important;
    }
    
    /* Scrollbar */
    ::-webkit-scrollbar {
        width: 10px;
        height: 10px;
    }
    
    ::-webkit-scrollbar-track {
        background: rgba(255, 255, 255, 0.05);
        border-radius: 10px;
    }
    
    ::-webkit-scrollbar-thumb {
        background: rgba(255, 255, 255, 0.2);
        border-radius: 10px;
    }
    
    ::-webkit-scrollbar-thumb:hover {
        background: rgba(0, 122, 255, 0.5);
    }
    
    /* Alert/Info boxes */
    .stAlert {
        background: rgba(255, 255, 255, 0.05) !important;
        border: 1px solid rgba(255, 255, 255, 0.1) !important;
        border-radius: 16px !important;
        backdrop-filter: blur(10px);
    }
    
    /* Text colors */
    p, span, div, label {
        color: #f5f5f7 !important;
    }
    
    h1, h2, h3, h4, h5, h6 {
        color: #ffffff !important;
    }
    
    /* Sidebar */
    [data-testid="stSidebar"] {
        background: rgba(0, 0, 0, 0.8) !important;
        backdrop-filter: blur(40px);
        border-right: 1px solid rgba(255, 255, 255, 0.1);
    }
    
    [data-testid="stSidebar"] .stMarkdown {
        color: #f5f5f7 !important;
    }
    </style>
""", unsafe_allow_html=True)

# Initialize session state
if "messages" not in st.session_state:
    st.session_state.messages = []
if "vectorstore" not in st.session_state:
    st.session_state.vectorstore = None
if "retriever" not in st.session_state:
    st.session_state.retriever = None
if "llm" not in st.session_state:
    st.session_state.llm = None
if "processed_files" not in st.session_state:
    st.session_state.processed_files = []
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "show_chat" not in st.session_state:
    st.session_state.show_chat = False

# PDF Generation Class
class QuestionPaperPDF(FPDF):
    def __init__(self):
        super().__init__()
        self.add_page()
        
    def header(self):
        self.set_font('Arial', 'B', 16)
        self.cell(0, 10, 'IIIT Sri City', 0, 1, 'C')
        self.set_font('Arial', 'B', 14)
        self.cell(0, 10, 'Question Paper', 0, 1, 'C')
        self.ln(5)
    
    def add_question_paper_content(self, content):
        self.set_font('Arial', '', 11)
        content = content.encode('latin-1', 'replace').decode('latin-1')
        self.multi_cell(0, 6, content)

def detect_question_paper_request(prompt):
    keywords = [
        'generate question paper', 'create question paper', 'make question paper',
        'generate questions', 'create exam', 'make test', 'question paper',
        'generate test', 'create test paper', 'quiz paper', 'exam paper'
    ]
    prompt_lower = prompt.lower()
    return any(keyword in prompt_lower for keyword in keywords)

def generate_question_paper_pdf(content, filename="question_paper.pdf"):
    pdf = QuestionPaperPDF()
    pdf.set_author('IIITSC Study-Buddy')
    pdf.set_title('Generated Question Paper')
    pdf.set_font('Arial', 'I', 10)
    pdf.cell(0, 10, f'Generated on: {datetime.now().strftime("%B %d, %Y")}', 0, 1, 'R')
    pdf.ln(5)
    pdf.add_question_paper_content(content)
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
    pdf.output(temp_file.name)
    return temp_file.name

def generate_question_paper_docx(content, filename="question_paper.docx"):
    """Generate a professionally formatted DOCX file for question papers"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add header - Institution name
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run('IIIT Sri City')
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 122, 255)
    
    # Add title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Question Paper')
    run.font.size = Pt(16)
    run.font.bold = True
    
    # Add date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = date_para.add_run(f'Generated on: {datetime.now().strftime("%B %d, %Y")}')
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Add a line separator
    doc.add_paragraph('_' * 80)
    
    # Process and add content
    lines = content.split('\n')
    for line in lines:
        if line.strip():
            para = doc.add_paragraph()
            
            # Detect headings (lines that are all caps or start with specific patterns)
            if line.strip().isupper() or line.strip().startswith('SECTION') or line.strip().startswith('PART'):
                run = para.add_run(line)
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Detect question numbers (lines starting with digits followed by . or ))
            elif re.match(r'^\d+[\.)]\s', line.strip()):
                run = para.add_run(line)
                run.font.size = Pt(12)
                run.font.bold = True
                para.space_before = Pt(6)
                para.space_after = Pt(3)
            
            # Detect sub-parts (a), b), i), ii), etc.)
            elif re.match(r'^[a-z\)ivx]+[\)]\s', line.strip().lower()):
                para.paragraph_format.left_indent = Inches(0.5)
                run = para.add_run(line)
                run.font.size = Pt(11)
            
            # Regular content
            else:
                run = para.add_run(line)
                run.font.size = Pt(11)
    
    # Add footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('Generated by Study-Buddy AI Assistant')
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save to temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    return temp_file.name

def format_chat_history(chat_history):
    formatted = []
    for msg in chat_history:
        if isinstance(msg, HumanMessage):
            formatted.append(f"Human: {msg.content}")
        elif isinstance(msg, AIMessage):
            formatted.append(f"Assistant: {msg.content}")
    return "\n".join(formatted)

def format_docs(docs):
    return "\n\n".join(doc.page_content for doc in docs)

# Hero Section
st.markdown("""
    <div class="hero-section">
        <div class="hero-title">Study-Buddy</div>
        <div class="hero-subtitle">AI-powered learning assistant that thinks like you do.</div>
    </div>
""", unsafe_allow_html=True)

# Feature Cards Section
if not st.session_state.show_chat:
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("""
            <div class="feature-card">
                <span class="feature-icon">🧠</span>
                <div class="feature-title">Smart RAG</div>
                <div class="feature-desc">Upload your documents and get instant, context-aware answers powered by advanced AI.</div>
            </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
            <div class="feature-card">
                <span class="feature-icon">📝</span>
                <div class="feature-title">Generate Papers</div>
                <div class="feature-desc">Create professional question papers in seconds with downloadable PDF format.</div>
            </div>
        """, unsafe_allow_html=True)
    
    with col3:
        st.markdown("""
            <div class="feature-card">
                <span class="feature-icon">💬</span>
                <div class="feature-title">Conversational</div>
                <div class="feature-desc">Natural dialogue with memory - it remembers your conversation context.</div>
            </div>
        """, unsafe_allow_html=True)

# API Key and Upload Section
st.markdown("<br>", unsafe_allow_html=True)

# Status indicator
status_col1, status_col2 = st.columns([3, 1])
with status_col2:
    if st.session_state.vectorstore and st.session_state.llm:
        st.markdown('<span class="status-badge status-active">● System Active</span>', unsafe_allow_html=True)
    else:
        st.markdown('<span class="status-badge status-inactive">● Configure Below</span>', unsafe_allow_html=True)

# Configuration Section
with st.expander("⚙️ Configuration", expanded=not st.session_state.llm):
    api_col1, api_col2 = st.columns([2, 1])
    
    with api_col1:
        if not os.getenv("GROQ_API_KEY"):
            api_key = st.text_input(
                "🔑 Groq API Key",
                type="password",
                help="Get your free API key from console.groq.com",
                placeholder="Enter your Groq API key..."
            )
            if api_key:
                os.environ["GROQ_API_KEY"] = api_key
                if not st.session_state.llm:
                    st.session_state.llm = ChatGroq(
                        model="llama-3.3-70b-versatile",
                        groq_api_key=api_key,
                        temperature=0.7
                    )
                    st.success("✓ API Key configured successfully!")

# Upload Section
st.markdown("""
    <div class="upload-section">
        <h2 style="color: #ffffff; margin-bottom: 1rem;">📚 Upload Your Documents</h2>
        <p style="color: #86868b; font-size: 1.1rem;">Drop your PDFs or text files to get started</p>
    </div>
""", unsafe_allow_html=True)

uploaded_files = st.file_uploader(
    "Choose files",
    type=["pdf", "txt"],
    accept_multiple_files=True,
    label_visibility="collapsed"
)

if uploaded_files:
    process_col1, process_col2, process_col3 = st.columns([1, 2, 1])
    with process_col2:
        if st.button("🚀 Process Documents", type="primary", use_container_width=True):
            if not os.getenv("GROQ_API_KEY"):
                st.error("⚠️ Please enter your Groq API Key first!")
            else:
                with st.spinner("Processing your documents..."):
                    try:
                        all_documents = []
                        
                        for uploaded_file in uploaded_files:
                            with tempfile.NamedTemporaryFile(delete=False, suffix=Path(uploaded_file.name).suffix) as tmp_file:
                                tmp_file.write(uploaded_file.getvalue())
                                tmp_path = tmp_file.name
                            
                            if uploaded_file.name.endswith('.pdf'):
                                loader = PyPDFLoader(tmp_path)
                            else:
                                loader = TextLoader(tmp_path)
                            
                            documents = loader.load()
                            all_documents.extend(documents)
                            os.unlink(tmp_path)
                        
                        text_splitter = RecursiveCharacterTextSplitter(
                            chunk_size=3000,
                            chunk_overlap=200,
                            length_function=len
                        )
                        splits = text_splitter.split_documents(all_documents)
                        
                        embeddings = HuggingFaceEmbeddings(
                            model_name="sentence-transformers/all-MiniLM-L6-v2"
                        )
                        
                        st.session_state.vectorstore = FAISS.from_documents(
                            documents=splits,
                            embedding=embeddings
                        )
                        
                        st.session_state.retriever = st.session_state.vectorstore.as_retriever(
                            search_kwargs={"k": 4}
                        )
                        
                        if not st.session_state.llm:
                            st.session_state.llm = ChatGroq(
                                model="llama-3.3-70b-versatile",
                                groq_api_key=os.getenv("GROQ_API_KEY"),
                                temperature=0.7
                            )
                        
                        st.session_state.processed_files = [f.name for f in uploaded_files]
                        st.session_state.show_chat = True
                        st.success(f"✅ Successfully processed {len(uploaded_files)} documents!")
                        st.rerun()
                        
                    except Exception as e:
                        st.error(f"❌ Error: {str(e)}")

# Show processed files
if st.session_state.processed_files:
    with st.expander("📄 Processed Documents"):
        for i, filename in enumerate(st.session_state.processed_files, 1):
            st.markdown(f"**{i}.** {filename}")

# Chat Interface
if st.session_state.show_chat or st.session_state.messages:
    st.markdown("<br><br>", unsafe_allow_html=True)
    st.markdown('<div class="chat-container">', unsafe_allow_html=True)
    
    # Display chat messages
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
    
    st.markdown('</div>', unsafe_allow_html=True)

# Chat Input
if prompt := st.chat_input("Ask anything or request 'generate question paper on [topic]'..."):
    if not st.session_state.llm and os.getenv("GROQ_API_KEY"):
        st.session_state.llm = ChatGroq(
            model="llama-3.3-70b-versatile",
            groq_api_key=os.getenv("GROQ_API_KEY"),
            temperature=0.7
        )
    
    if not st.session_state.llm:
        st.warning("⚠️ Please configure your Groq API key first!")
    else:
        st.session_state.show_chat = True
        st.session_state.messages.append({"role": "user", "content": prompt})
        
        with st.chat_message("user"):
            st.markdown(prompt)
        
        with st.chat_message("assistant"):
            with st.spinner("Thinking..."):
                try:
                    is_qp_request = detect_question_paper_request(prompt)
                    
                    if is_qp_request:
                        qp_template = """You are an expert educator. Generate a comprehensive question paper based on the user's request.
Include various types of questions (MCQ, short answer, long answer) with clear instructions.
Format it professionally with proper numbering and sections.

User Request: {question}

Context from documents (if available):
{context}

Generate a well-structured question paper:"""
                        
                        context = ""
                        if st.session_state.retriever:
                            retrieved_docs = st.session_state.retriever.invoke(prompt)
                            context = format_docs(retrieved_docs)
                        
                        qp_prompt = ChatPromptTemplate.from_template(qp_template)
                        
                        chain = (
                            {
                                "context": lambda x: context,
                                "question": RunnablePassthrough()
                            }
                            | qp_prompt
                            | st.session_state.llm
                            | StrOutputParser()
                        )
                        
                        answer = chain.invoke(prompt)
                        st.markdown(answer)
                        
                        # Generate both PDF and DOCX
                        pdf_path = generate_question_paper_pdf(answer)
                        docx_path = generate_question_paper_docx(answer)
                        
                        # Create two download buttons side by side
                        download_col1, download_col2 = st.columns(2)
                        
                        with download_col1:
                            with open(pdf_path, "rb") as pdf_file:
                                st.download_button(
                                    label="📥 Download PDF",
                                    data=pdf_file,
                                    file_name=f"question_paper_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                                    mime="application/pdf",
                                    use_container_width=True
                                )
                        
                        with download_col2:
                            with open(docx_path, "rb") as docx_file:
                                st.download_button(
                                    label="📄 Download DOCX",
                                    data=docx_file,
                                    file_name=f"question_paper_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    type="primary",
                                    use_container_width=True
                                )
                        
                        # Clean up temporary files
                        os.unlink(pdf_path)
                        os.unlink(docx_path)
                        
                    else:
                        if st.session_state.retriever:
                            template = """You are a helpful AI assistant for IIIT Sri City students. 
Answer the question using the context from documents if relevant, otherwise use your general knowledge.
Be friendly, professional, and detailed.

Context from documents:
{context}

Chat History:
{chat_history}

Current Question: {question}

Answer:"""
                            
                            prompt_template = ChatPromptTemplate.from_template(template)
                            retrieved_docs = st.session_state.retriever.invoke(prompt)
                            context = format_docs(retrieved_docs)
                            chat_history_text = format_chat_history(st.session_state.chat_history[-6:])
                            
                            rag_chain = (
                                {
                                    "context": lambda x: context,
                                    "chat_history": lambda x: chat_history_text,
                                    "question": RunnablePassthrough()
                                }
                                | prompt_template
                                | st.session_state.llm
                                | StrOutputParser()
                            )
                            
                            answer = rag_chain.invoke(prompt)
                            
                            if retrieved_docs:
                                with st.expander("📚 View Sources"):
                                    for i, doc in enumerate(retrieved_docs, 1):
                                        st.markdown(f"**Source {i}:**")
                                        st.text(doc.page_content[:300] + "...")
                                        if hasattr(doc, 'metadata') and 'source' in doc.metadata:
                                            st.caption(f"From: {Path(doc.metadata['source']).name}")
                                        st.markdown("---")
                        else:
                            answer = st.session_state.llm.invoke([HumanMessage(content=prompt)]).content
                        
                        st.markdown(answer)
                    
                    st.session_state.chat_history.append(HumanMessage(content=prompt))
                    st.session_state.chat_history.append(AIMessage(content=answer))
                    st.session_state.messages.append({"role": "assistant", "content": answer})
                    
                except Exception as e:
                    error_msg = f"❌ Error: {str(e)}"
                    st.error(error_msg)
                    st.session_state.messages.append({"role": "assistant", "content": error_msg})

# Footer
st.markdown("<br><br>", unsafe_allow_html=True)
st.markdown("""
    <div style='text-align: center; padding: 3rem 0; color: #86868b;'>
        <p style='font-size: 0.95rem; margin-bottom: 0.5rem;'>Designed for students. Built with passion.</p>
        <p style='font-size: 0.85rem;'>IIIT Sri City © 2026</p>
    </div>
""", unsafe_allow_html=True)